from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        shading = c._element.get_or_add_tcPr()
        sh = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
        shading.append(sh)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def verdict(text, color='green'):
    p = doc.add_paragraph()
    r = p.add_run(f'ВЕРДИКТ: {text}')
    r.bold = True
    r.font.size = Pt(10)
    colors = {'green': RGBColor(0,128,0), 'red': RGBColor(200,0,0), 'orange': RGBColor(200,120,0)}
    r.font.color.rgb = colors.get(color, RGBColor(0,0,0))

# ================================================================
# TITLE PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('НЕЗАВИСИМЫЙ АУДИТ\nМЕТОДИК РАСЧЁТА')
r.bold = True
r.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Платформа АЕМР — Аналитика закупок\nЕлизовский муниципальный район')
r.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Верификация: СВОД ТД-ПМ vs Код АЕМР vs Независимый расчёт')
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Апрель 2026')
r.font.size = Pt(12)

doc.add_page_break()

# ================================================================
# TOC
# ================================================================
doc.add_heading('Содержание', level=1)
toc_items = [
    '1. Введение и методология аудита',
    '2. Структура исходных данных',
    '3. Количество закупок (plan_count / fact_count)',
    '4. Бюджетные суммы (ФБ / КБ / МБ / Итого)',
    '5. Отклонение (deviation)',
    '6. Процент исполнения (execution %)',
    '7. Экономия (economy)',
    '8. Разделение КП / ЕП',
    '9. Годовые итоги и агрегация',
    '10. Помесячная динамика (ШДЮ)',
    '11. Сводная таблица соответствия',
    '12. Выводы и рекомендации',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ================================================================
# 1. INTRODUCTION
# ================================================================
doc.add_heading('1. Введение и методология аудита', level=1)

doc.add_paragraph(
    'Настоящий документ представляет собой независимую верификацию всех методик расчёта, '
    'используемых в платформе АЕМР для аналитики закупок Елизовского муниципального района. '
    'Для каждой категории метрик рассматриваются три перспективы:'
)

items = [
    ('СВОД ТД-ПМ', 'формулы Google Sheets (COUNTIFS/SUMIFS), применяемые в официальной сводной таблице'),
    ('Код АЕМР', 'реализация в TypeScript (recalculate.ts, orchestrator.ts), исполняемая на сервере'),
    ('Математически верный метод', 'определяется на основе экономической логики закупочной деятельности, '
     'требований 44-ФЗ и общепринятых практик финансовой аналитики'),
]
for title, desc in items:
    p = doc.add_paragraph()
    r = p.add_run(f'{title} — ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('Критерии оценки', level=2)
doc.add_paragraph(
    'Каждая методика оценивается по следующим критериям: '
    'корректность формулы, полнота охвата данных, устойчивость к краевым случаям '
    '(пустые ячейки, заглушки, многолетние данные), соответствие экономическому смыслу показателя.'
)

doc.add_heading('Источники данных', level=2)
add_table(
    ['Источник', 'Описание', 'Объём'],
    [
        ['СВОД ТД-ПМ', 'Сводная таблица, 24 столбца (A-X), 279 строк', '2 блока (КП/ЕП) x 9 ГРБС'],
        ['Листы управлений (8 шт.)', 'Построчные данные, 33 столбца (A-AG)', '~50-120 строк на управление'],
        ['ШДЮ', 'Помесячная динамика, 21 столбец', '~300 строк (8 ГРБС x ~35 строк)'],
    ],
    col_widths=[4, 8, 5]
)

doc.add_page_break()

# ================================================================
# 2. DATA STRUCTURE
# ================================================================
doc.add_heading('2. Структура исходных данных', level=1)

doc.add_heading('2.1. Листы управлений (department sheets)', level=2)
doc.add_paragraph('Каждый лист управления содержит построчные записи о закупках. Ключевые столбцы:')

add_table(
    ['Столбец', 'Индекс', 'Название', 'Описание'],
    [
        ['A', '0', 'ID', 'Порядковый номер'],
        ['B', '1', 'Реестровый номер', 'Номер в реестре закупок'],
        ['C', '2', 'Подведомственный', 'Наименование организации'],
        ['F', '5', 'Вид деятельности', 'Программное мероприятие / Текущая деятельность'],
        ['H', '7', 'ФБ план', 'Федеральный бюджет (план), руб.'],
        ['I', '8', 'КБ план', 'Краевой бюджет (план), руб.'],
        ['J', '9', 'МБ план', 'Муниципальный бюджет (план), руб.'],
        ['K', '10', 'Итого план', 'Сумма H+I+J или формула'],
        ['L', '11', 'Способ', 'ЭА / ЕП / ЭК / ЭЗК'],
        ['O', '14', 'Квартал плана', '1, 2, 3 или 4'],
        ['P', '15', 'Год плана', '2025, 2026 и т.д.'],
        ['Q', '16', 'Дата факта', 'Дата контракта или заглушка (Х, -)'],
        ['R', '17', 'Квартал факта', '1, 2, 3 или 4'],
        ['V', '21', 'ФБ факт', 'Федеральный бюджет (факт), руб.'],
        ['W', '22', 'КБ факт', 'Краевой бюджет (факт), руб.'],
        ['X', '23', 'МБ факт', 'Муниципальный бюджет (факт), руб.'],
        ['Y', '24', 'Итого факт', 'Сумма V+W+X или формула'],
        ['AD', '29', 'Признак экономии', '"да" / "нет" / пусто'],
    ],
    col_widths=[2, 2, 4, 9]
)

doc.add_heading('2.2. СВОД ТД-ПМ (summary sheet)', level=2)
doc.add_paragraph(
    'Сводная таблица содержит агрегированные показатели, рассчитанные формулами COUNTIFS/SUMIFS. '
    'Структура: два блока (КП - конкурентные процедуры, строки 3-14; ЕП - единственный поставщик, строки 15-26), '
    'внутри каждого - строки по кварталам (Q1-Q4) и итоговые строки за год.'
)

add_table(
    ['Столбец СВОД', 'Буква', 'Содержание'],
    [
        ['A', 'A', 'ГРБС (наименование управления)'],
        ['B', 'B', 'Квартал (1-4)'],
        ['C', 'C', 'Год (2025, 2026)'],
        ['D', 'D', 'План количество (COUNTIFS)'],
        ['E', 'E', 'Факт количество (COUNTIFS)'],
        ['F', 'F', 'Отклонение (E-D)'],
        ['G', 'G', 'Выполнено % (E/D)'],
        ['H-K', 'H-K', 'Бюджеты план (ФБ, КБ, МБ, Итого) - SUMIFS'],
        ['L-O', 'L-O', 'Бюджеты факт (ФБ, КБ, МБ, Итого) - SUMIFS'],
        ['P', 'P', 'Отклонение сумм (O-K)'],
        ['Q', 'Q', 'Потрачено % (O/K)'],
        ['R-U', 'R-U', 'Экономия (ФБ, КБ, МБ, Итого)'],
    ],
    col_widths=[3, 2, 12]
)

doc.add_page_break()

# ================================================================
# 3. PLAN/FACT COUNT
# ================================================================
doc.add_heading('3. Количество закупок (plan_count / fact_count)', level=1)

doc.add_heading('3.1. Как считает СВОД', level=2)
doc.add_paragraph('Формула plan_count для КП, квартал Q1, управление УЭР:')
p = doc.add_paragraph()
r = p.add_run('=COUNTIFS(УЭР!L:L;"<>ЕП"; УЭР!O:O;1; УЭР!P:P;2026)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph('Условия фильтрации:')
items_count = [
    'L<>ЕП - метод закупки не равен "ЕП" (т.е. конкурентные: ЭА, ЭК, ЭЗК)',
    'O=1 - квартал плана равен 1 (Q1)',
    'P=2026 - год плана равен целевому году',
]
for item in items_count:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Формула fact_count для КП, квартал Q1:')
p = doc.add_paragraph()
r = p.add_run('=COUNTIFS(УЭР!L:L;"<>ЕП"; УЭР!O:O;1; УЭР!P:P;2026; УЭР!Q:Q;"<>Х"; УЭР!Q:Q;"<>X"; УЭР!Q:Q;"<>")')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph('Дополнительные условия для факта:')
items_fact = [
    'Q<>Х - дата факта не равна кириллической заглушке "Х"',
    'Q<>X - дата факта не равна латинской заглушке "X"',
    'Q<>"" - дата факта не пуста',
]
for item in items_fact:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Для ЕП аналогично, но L="ЕП" вместо L<>"ЕП".')

doc.add_heading('3.2. Как считает код АЕМР', level=2)
doc.add_paragraph('Файл: packages/core/src/pipeline/recalculate.ts, функция recalculateFromRows()')
doc.add_paragraph('Алгоритм построчной обработки:')

steps = [
    'Пропуск строк без ID (столбец A) и без предмета (столбец G)',
    'Пропуск сводных строк (начинаются с "итого", "всего", "справочно")',
    'Классификация строки: score >= 3 из {метод(+3), тип(+2), суммы>0(+2), дата(+1), ID/предмет(+1)}',
    'Фильтр по году: если targetYear задан, пропуск строк где P != targetYear',
    'Определение метода: ЭА/ЭК/ЭЗК -> competitive (КП), ЕП -> sole (ЕП)',
    'Plan count: ++planCount для квартала из столбца O (если planQ != null)',
    'Fact detection: столбец Q непуст И не является заглушкой (Х, X, -, -, -, н/д, нет, не определена)',
    'Fact count: если hasFact == true -> ++factCount в соответствующем квартале',
    'Year plan: только если planQ не пуст (COUNTIFS гейт на столбец O)',
    'Year fact: сумма квартальных fact_count + прямое добавление для строк без planQ',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('3.3. Независимая верификация', level=2)

add_table(
    ['Аспект', 'СВОД', 'Код АЕМР', 'Верный метод'],
    [
        ['Фильтр метода', 'L<>"ЕП" / L="ЕП"', 'COMPETITIVE_METHODS.has(method) / method === "ЕП"',
         'Эквивалентно. Код точнее: явный перечень {ЭА, ЭК, ЭЗК}'],
        ['Фильтр квартала', 'O=N (COUNTIFS)', 'getQuarterKey(row[O])', 'Эквивалентно'],
        ['Фильтр года', 'P=2026 (COUNTIFS)', 'targetYear filter на col P',
         'Эквивалентно. Критично для многолетних листов'],
        ['Факт: заглушки', 'Q<>"Х"; Q<>"X"; Q<>""', 'PLACEHOLDERS set (7 значений)',
         'Код строже: ловит "-", "н/д", "нет". СВОД может пропустить "-"'],
        ['Plan count гейт', 'COUNTIFS на O (квартал)', 'if (planQ) year.planCount++',
         'Эквивалентно после фикса. Без фикса был inflation'],
        ['Классификация', 'Нет (считает все строки)', 'Score-based (>= 3)',
         'Код корректнее: фильтрует пустые/служебные строки'],
    ],
    col_widths=[3, 4, 4.5, 5.5]
)

verdict(
    'Plan/fact count: методики СВОД и АЕМР эквивалентны после внесённых исправлений. '
    'Код АЕМР строже в обработке заглушек (7 паттернов vs 2 в СВОД), что является преимуществом. '
    'Математически верный метод соответствует реализации АЕМР.',
    'green'
)

doc.add_page_break()

# ================================================================
# 4. BUDGETS
# ================================================================
doc.add_heading('4. Бюджетные суммы (ФБ / КБ / МБ / Итого)', level=1)

doc.add_heading('4.1. Как считает СВОД', level=2)
doc.add_paragraph('Формула plan_fb для КП, Q1, УЭР:')
p = doc.add_paragraph()
r = p.add_run('=SUMIFS(УЭР!H:H; УЭР!L:L;"<>ЕП"; УЭР!O:O;1; УЭР!P:P;2026)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph(
    'Логика идентична plan_count, но вместо COUNTIFS используется SUMIFS. '
    'Для каждого бюджета (ФБ=H, КБ=I, МБ=J) - отдельная формула с теми же условиями фильтрации. '
    'Итого (K) = SUMIFS на столбец K, или формула =H+I+J.'
)

doc.add_paragraph('Для фактических бюджетов (fact_fb/kb/mb):')
p = doc.add_paragraph()
r = p.add_run('=SUMIFS(УЭР!V:V; УЭР!L:L;"<>ЕП"; УЭР!O:O;1; УЭР!P:P;2026; УЭР!Q:Q;"<>Х"; ...)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph(
    'Фактические бюджеты (V=ФБ, W=КБ, X=МБ, Y=Итого) суммируются только для строк '
    'с заполненной датой факта (те же условия Q<>"Х", Q<>"X", Q<>"").'
)

doc.add_heading('4.2. Как считает код АЕМР', level=2)
doc.add_paragraph('Файл: packages/core/src/pipeline/recalculate.ts')
doc.add_paragraph('Бюджеты накапливаются в цикле по строкам:')

budget_steps = [
    'Plan budgets (H, I, J): суммируются для строк с непустым planQ, раздельно для competitive и ep',
    'Plan total: planTotalFor(K, H, I, J) - берёт K если != 0, иначе H+I+J',
    'Fact budgets (V, W, X): суммируются только если hasFact == true',
    'Fact total: factTotalFor(Y, V, W, X) - берёт Y если != 0, иначе V+W+X',
    'Per-method split: competitive.planFB += hVal (если isCompetitive), ep.planFB += hVal (если isEP)',
    'Year totals: plan из if(planQ), fact из суммы кварталов',
]
for step in budget_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('4.3. Независимая верификация', level=2)

add_table(
    ['Аспект', 'СВОД', 'Код АЕМР', 'Верный метод'],
    [
        ['Plan budget гейт', 'SUMIFS на O (квартал)', 'if (planQ) q.planFB += hVal',
         'Эквивалентно. Plan суммы привязаны к кварталу'],
        ['Fact budget гейт', 'SUMIFS + Q<>Х/X/""', 'if (hasFact) q.factFB += vVal',
         'Эквивалентно. Код строже (7 заглушек)'],
        ['Итого план', 'SUMIFS на K', 'planTotalFor: K || H+I+J',
         'Код корректнее: fallback на сумму компонент при K=0'],
        ['Итого факт', 'SUMIFS на Y', 'factTotalFor: Y || V+W+X',
         'Код корректнее: fallback при Y=0'],
        ['КП/ЕП разделение', 'Отдельные SUMIFS по L', 'if(isCompetitive)/if(isEP) split',
         'Эквивалентно. Оба используют метод как ключ'],
    ],
    col_widths=[3, 4, 4.5, 5.5]
)

verdict(
    'Бюджетные суммы: методики эквивалентны. Код АЕМР имеет преимущество: '
    'fallback на сумму компонент (H+I+J) при пустом итоге (K), '
    'что повышает устойчивость к неполным данным.',
    'green'
)

doc.add_page_break()

# ================================================================
# 5. DEVIATION
# ================================================================
doc.add_heading('5. Отклонение (deviation)', level=1)

doc.add_heading('5.1. Как считает СВОД', level=2)
doc.add_paragraph('Столбец F в СВОД ТД-ПМ:')
p = doc.add_paragraph()
r = p.add_run('F = E - D  (fact_count - plan_count)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.bold = True

doc.add_paragraph(
    'Отклонение - это разность между количеством выполненных (факт) и запланированных закупок. '
    'Отрицательное значение означает недовыполнение плана, положительное - перевыполнение.'
)

doc.add_heading('5.2. Как считает код АЕМР (после исправления)', level=2)
doc.add_paragraph('Файл: packages/core/src/pipeline/orchestrator.ts, mergeSummaryMetrics()')
p = doc.add_paragraph()
r = p.add_run('const kpDeviation = kpFact - kpCount;')
r.font.name = 'Consolas'
r.font.size = Pt(10)

doc.add_paragraph(
    'До исправления код использовал формулу plan - fact (kpCount - kpFact), '
    'что давало инвертированный знак. Исправлено на fact - plan для соответствия СВОД.'
)

doc.add_heading('5.3. Независимая верификация', level=2)
doc.add_paragraph(
    'В экономике закупок отклонение = факт - план является стандартом. '
    'Отрицательное значение сигнализирует о невыполнении, положительное - о сверхплановых закупках. '
    'Формула СВОД (E-D) и исправленная формула кода (kpFact - kpCount) корректны.'
)

verdict(
    'Отклонение: после исправления знака формулы код полностью соответствует СВОД и математическому стандарту.',
    'green'
)

# ================================================================
# 6. EXECUTION %
# ================================================================
doc.add_heading('6. Процент исполнения (execution %)', level=1)

doc.add_heading('6.1. Как считает СВОД', level=2)
doc.add_paragraph('Столбец G в СВОД ТД-ПМ:')
p = doc.add_paragraph()
r = p.add_run('G = E / D  (fact_count / plan_count)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.bold = True
doc.add_paragraph('Столбец Q (по суммам):')
p = doc.add_paragraph()
r = p.add_run('Q = O / K  (fact_total / plan_total)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.bold = True

doc.add_paragraph(
    'Исполнение по количеству: доля выполненных от запланированных. '
    'Исполнение по суммам: доля освоенных бюджетов от запланированных. '
    'При D=0 или K=0 ячейка содержит ошибку деления или 0.'
)

doc.add_heading('6.2. Как считает код АЕМР', level=2)
doc.add_paragraph('Файл: recalculate.ts, функция pct():')
p = doc.add_paragraph()
r = p.add_run('function pct(part, total) { return total > 0 ? +(part / total).toFixed(6) : 0; }')
r.font.name = 'Consolas'
r.font.size = Pt(10)

doc.add_paragraph(
    'Код использует: executionPct = pct(factTotal, planTotal). '
    'При planTotal <= 0 возвращает 0, что безопаснее деления СВОД (возможен #DIV/0!).'
)

doc.add_heading('6.3. Независимая верификация', level=2)

add_table(
    ['Аспект', 'СВОД', 'Код АЕМР', 'Верный метод'],
    [
        ['Формула', 'fact / plan', 'pct(fact, plan)', 'Эквивалентно'],
        ['Деление на 0', 'Ошибка #DIV/0!', 'Возвращает 0', 'Код корректнее: 0 вместо ошибки'],
        ['По количеству', 'E/D', 'fact_count/plan_count', 'Эквивалентно'],
        ['По суммам', 'O/K', 'factTotal/planTotal', 'Эквивалентно'],
        ['Шкала', '0-100%', '0.0-1.0 (decimal)', 'Вопрос представления, не логики'],
    ],
    col_widths=[3, 4, 4.5, 5.5]
)

verdict(
    'Процент исполнения: методики математически эквивалентны. '
    'Код безопаснее: обрабатывает деление на 0 без ошибки.',
    'green'
)

doc.add_page_break()

# ================================================================
# 7. ECONOMY
# ================================================================
doc.add_heading('7. Экономия (economy)', level=1)

doc.add_heading('7.1. Как считает СВОД', level=2)
doc.add_paragraph('Столбцы R-U в СВОД ТД-ПМ. Экономия по МБ (столбец T):')
p = doc.add_paragraph()
r = p.add_run('=SUMIFS(УЭР!J:J; УЭР!AD:AD;"да"; ...) - SUMIFS(УЭР!X:X; УЭР!AD:AD;"да"; ...)')
r.font.name = 'Consolas'
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph('Условия экономии в СВОД:')
items_eco = [
    'AD="да" - признак экономии установлен (столбец AD содержит "да")',
    'L<>"ЕП" / L="ЕП" - разделение по методу закупки (КП/ЕП)',
    'O=квартал, P=год - стандартные фильтры периода',
    'Экономия = Plan_MB - Fact_MB (J - X), гейтирована флагом AD',
]
for item in items_eco:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.2. Как считает код АЕМР', level=2)
doc.add_paragraph('Файл: recalculate.ts, строки 568-587:')

eco_steps = [
    'Условие: hasFact == true И J (MB план) > 0',
    'Формула: eco = J - X (plan_mb - fact_mb)',
    'Фильтр: eco > 0 (только положительная экономия)',
    'НЕ проверяет столбец AD ("да"/"нет")',
    'Накопление: year.economyMB += eco, quarters[planQ].economyMB += eco',
    'Per-method split: competitive.economyMB / ep.economyMB',
]
for step in eco_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('7.3. Ключевое расхождение: флаг AD', level=2)

p = doc.add_paragraph()
r = p.add_run('КРИТИЧЕСКОЕ РАЗЛИЧИЕ: ')
r.bold = True
r.font.color.rgb = RGBColor(200, 0, 0)
p.add_run(
    'СВОД гейтирует экономию флагом AD="да", код АЕМР - нет. '
    'Код считает экономию математически (J-X>0 при наличии факта), '
    'СВОД полагается на ручной флаг, установленный оператором.'
)

add_table(
    ['Аспект', 'СВОД', 'Код АЕМР', 'Верный метод'],
    [
        ['Условие экономии', 'AD="да" (ручной флаг)', 'J > 0 AND J-X > 0 AND hasFact',
         'Гибридный: J-X>0 как базис + AD для подтверждения'],
        ['Направление', 'Plan - Fact (J - X)', 'Plan - Fact (J - X)', 'Корректно: экономия = снижение цены'],
        ['Отрицательная экономия', 'Не считается (AD не ставят)', 'Отсекается eco > 0',
         'Корректно: перерасход это другой показатель'],
        ['Без факта', 'Не считается (нет Q)', 'Не считается (hasFact)', 'Корректно: нет факта = нет экономии'],
        ['AD конфликт', 'Не детектирует', 'conflicts++ при несовпадении',
         'Код лучше: обнаруживает ошибки оператора'],
    ],
    col_widths=[3, 4, 4.5, 5.5]
)

doc.add_heading('7.4. Экономическое обоснование', level=2)
doc.add_paragraph(
    'С точки зрения экономики закупок (44-ФЗ), экономия - это разница между начальной максимальной '
    'ценой контракта (НМЦК, план) и ценой заключённого контракта (факт) при условии, что контракт заключён. '
    'Математический метод (J-X>0 при наличии факта) точнее отражает реальность, чем ручной флаг AD, '
    'который может быть не установлен по ошибке оператора. Однако флаг AD может учитывать нюансы: '
    'например, снижение цены из-за изменения объёма, а не конкурентного торга.'
)

verdict(
    'Экономия: методики различаются. СВОД использует ручной флаг (AD="да"), код - математический расчёт (J-X>0). '
    'Оптимальный метод: гибридный - математический расчёт как базис, '
    'флаг AD как подтверждение, детекция конфликтов при несовпадении. '
    'Код АЕМР ближе к оптимальному методу.',
    'orange'
)

doc.add_page_break()

# ================================================================
# 8. KP/EP SPLIT
# ================================================================
doc.add_heading('8. Разделение КП / ЕП', level=1)

doc.add_heading('8.1. Определения', level=2)
add_table(
    ['Категория', 'Методы', 'Описание'],
    [
        ['КП (конкурентные)', 'ЭА, ЭК, ЭЗК', 'Электронный аукцион, электронный конкурс, запрос котировок'],
        ['ЕП (единственный)', 'ЕП', 'Закупка у единственного поставщика (44-ФЗ ст. 93)'],
    ],
    col_widths=[4, 3, 10]
)

doc.add_heading('8.2. Как разделяет СВОД', level=2)
doc.add_paragraph(
    'СВОД использует двойную фильтрацию в COUNTIFS/SUMIFS: '
    'для КП - L<>"ЕП" (всё кроме ЕП), для ЕП - L="ЕП". '
    'Это означает, что если появится новый метод (не ЭА/ЕП/ЭК/ЭЗК), '
    'СВОД отнесёт его к КП (по принципу исключения).'
)

doc.add_heading('8.3. Как разделяет код АЕМР', level=2)
doc.add_paragraph(
    'Код использует множество COMPETITIVE_METHODS = {ЭА, ЭК, ЭЗК} и проверку method === "ЕП". '
    'Строка с неизвестным методом не попадёт ни в КП, ни в ЕП (но войдёт в общие итоги). '
    'Это строже, чем СВОД: неизвестный метод не искажает ни одну категорию.'
)

doc.add_heading('8.4. Независимая верификация', level=2)

add_table(
    ['Аспект', 'СВОД', 'Код АЕМР', 'Верный метод'],
    [
        ['КП фильтр', 'L<>"ЕП" (всё кроме ЕП)', 'COMPETITIVE_METHODS.has(L)',
         'Код строже и точнее'],
        ['ЕП фильтр', 'L="ЕП"', 'method === "ЕП"', 'Эквивалентно'],
        ['Неизвестный метод', 'Попадает в КП', 'Не попадает ни в КП, ни в ЕП',
         'Код корректнее: не искажает данные'],
        ['Общие итоги', 'КП + ЕП = всё', 'КП + ЕП <= всего (возможен остаток)',
         'Код точнее: честный учёт'],
    ],
    col_widths=[3.5, 4, 4.5, 5]
)

verdict(
    'Разделение КП/ЕП: код АЕМР строже и точнее. СВОД группирует по исключению (L<>"ЕП"), '
    'что может включить мусорные данные в КП. На практике разница минимальна (нет неизвестных методов в текущих данных).',
    'green'
)

doc.add_page_break()

# ================================================================
# 9. YEAR TOTALS
# ================================================================
doc.add_heading('9. Годовые итоги и агрегация', level=1)

doc.add_heading('9.1. Как считает СВОД', level=2)
doc.add_paragraph(
    'Годовые итоги в СВОД - это SUM квартальных строк. '
    'Например, для УЭР КП: итого_plan = SUM(D_Q1:D_Q4) где D - столбец plan_count. '
    'Каждая квартальная строка рассчитана через COUNTIFS/SUMIFS. '
    'Таким образом, годовой итог = сумма квартальных COUNTIFS.'
)

doc.add_heading('9.2. Как считает код АЕМР', level=2)
doc.add_paragraph(
    'Год (plan): суммируется в цикле для строк с непустым planQ: if(planQ) year.planCount++. '
    'Год (fact): сумма квартальных factCount + прямые факты без planQ.'
)
p = doc.add_paragraph()
r = p.add_run(
    '// Year fact totals from quarter sums\n'
    'for (const qk of [q1..q4]) {\n'
    '  result.year.factCount += result.quarters[qk].factCount;\n'
    '}\n'
    '// + direct facts without planQ (rare edge case)'
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_heading('9.3. Критический фикс: фильтрация по году', level=2)
p = doc.add_paragraph()
r = p.add_run('ИСПРАВЛЕННЫЙ БАГ: ')
r.bold = True
r.font.color.rgb = RGBColor(200, 0, 0)
p.add_run(
    'До исправления код не фильтровал по столбцу P (год). Листы управлений содержат данные за 2025 и 2026 годы. '
    'СВОД фильтрует через COUNTIFS(P:P;2026). Без этого фильтра код считал все годы, '
    'давая inflation: 610 (код) vs 462 (СВОД). После добавления targetYear фильтра - совпадение.'
)

doc.add_heading('9.4. Агрегация по управлениям -> сводный уровень', level=2)
doc.add_paragraph(
    'Код: mergeSummaryMetrics() в orchestrator.ts суммирует все 8 управлений. '
    'СВОД: строки 7-14 (блок "ВСЕ") содержат COUNTIFS/SUMIFS по всем листам. '
    'Логика идентична: сумма department = total.'
)

verdict(
    'Годовые итоги: после добавления фильтра по году (targetYear) и гейта planQ - полное соответствие СВОД. '
    'Математически верный метод = сумма квартальных значений, что реализовано в обоих подходах.',
    'green'
)

doc.add_page_break()

# ================================================================
# 10. SHDYU
# ================================================================
doc.add_heading('10. Помесячная динамика (ШДЮ)', level=1)

doc.add_heading('10.1. Структура данных ШДЮ', level=2)
doc.add_paragraph(
    'ШДЮ (Ш-Д-Ю) - отдельная таблица помесячной динамики из файла СВОД_для_Google. '
    'Содержит те же показатели, что и СВОД ТД-ПМ, но в помесячном разрезе (12 месяцев вместо 4 кварталов).'
)

add_table(
    ['Столбец', 'Индекс', 'Содержание'],
    [
        ['A', '0', 'ГРБС (имя управления)'],
        ['B', '1', 'Заголовок блока (КП/ЕП)'],
        ['C', '2', 'Подзаголовок'],
        ['D', '3', 'Номер месяца (1-12)'],
        ['E', '4', 'Plan count'],
        ['F', '5', 'Fact count'],
        ['G', '6', 'Deviation'],
        ['H', '7', 'Execution %'],
        ['I-L', '8-11', 'Plan budgets (ФБ, КБ, МБ, Итого)'],
        ['M-P', '12-15', 'Fact budgets (ФБ, КБ, МБ, Итого)'],
        ['Q', '16', 'Budget deviation'],
        ['R', '17', 'Spent %'],
        ['S-V', '18-21', 'Economy (ФБ, КБ, МБ, Итого)'],
    ],
    col_widths=[2.5, 2, 12.5]
)

doc.add_heading('10.2. Использование в АЕМР', level=2)
doc.add_paragraph(
    'ШДЮ данные загружаются параллельно с основными листами (fetchSHDYUSheet в snapshot.ts), '
    'парсятся через parseSHDYUSheet() и сохраняются в snapshot.shdyuData. '
    'Используются для помесячной сверки (reconciliation) и тепловой карты на странице Recon.'
)

doc.add_heading('10.3. Верификация', level=2)
doc.add_paragraph(
    'ШДЮ - это независимый источник данных с той же логикой COUNTIFS/SUMIFS, но по месяцам. '
    'Квартальный итог в ШДЮ должен совпадать с квартальными данными СВОД ТД-ПМ: '
    'сумма(month1..month3) = Q1, сумма(month4..month6) = Q2 и т.д. '
    'Расхождения указывают на ошибку в одном из источников.'
)

verdict(
    'ШДЮ: данные используются как третий источник верификации. '
    'Помесячная гранулярность позволяет выявить аномалии внутри кварталов, '
    'недоступные при квартальной агрегации.',
    'green'
)

doc.add_page_break()

# ================================================================
# 11. SUMMARY TABLE
# ================================================================
doc.add_heading('11. Сводная таблица соответствия методик', level=1)

add_table(
    ['Показатель', 'СВОД = Код?', 'Кто точнее', 'Оценка'],
    [
        ['Plan count (КП)', 'Да', 'Эквивалентно', 'Полное соответствие'],
        ['Plan count (ЕП)', 'Да', 'Эквивалентно', 'Полное соответствие'],
        ['Fact count (КП)', 'Да', 'Код (7 заглушек)', 'Полное соответствие'],
        ['Fact count (ЕП)', 'Да', 'Код (7 заглушек)', 'Полное соответствие'],
        ['Отклонение', 'Да', 'Эквивалентно', 'После фикса знака'],
        ['Исполнение %', 'Да', 'Код (no DIV/0)', 'Полное соответствие'],
        ['ФБ план/факт', 'Да', 'Эквивалентно', 'Полное соответствие'],
        ['КБ план/факт', 'Да', 'Эквивалентно', 'Полное соответствие'],
        ['МБ план/факт', 'Да', 'Код (fallback K=0)', 'Полное соответствие'],
        ['Итого план/факт', 'Да', 'Код (fallback)', 'Полное соответствие'],
        ['Экономия', 'Нет', 'Код (математика)', 'Расхождение по гейту AD'],
        ['КП/ЕП разделение', '~Да', 'Код (явный перечень)', 'Практически эквивалентно'],
        ['Годовые итоги', 'Да', 'Эквивалентно', 'После фикса targetYear'],
        ['ШДЮ месячные', 'Н/П', 'Третий источник', 'Перекрёстная проверка'],
    ],
    col_widths=[3.5, 2.5, 4, 7]
)

doc.add_page_break()

# ================================================================
# 12. CONCLUSIONS
# ================================================================
doc.add_heading('12. Выводы и рекомендации', level=1)

doc.add_heading('12.1. Общая оценка', level=2)
doc.add_paragraph(
    'Проведённый аудит показал, что после внесённых исправлений методики расчёта кода АЕМР '
    'и формулы СВОД ТД-ПМ математически эквивалентны по 12 из 13 категорий показателей. '
    'Единственное значимое расхождение - в методике расчёта экономии (гейт AD vs математический расчёт).'
)

doc.add_heading('12.2. Исправленные баги', level=2)
bugs = [
    ('Year count inflation', 'Отсутствие фильтра по столбцу P (год). calc=610 vs СВОД=462. '
     'Фикс: добавлен targetYear параметр и фильтрация.'),
    ('Deviation sign', 'Инвертированный знак: plan-fact вместо fact-plan. '
     'Фикс: kpFact - kpCount.'),
    ('Fact count overcount', 'Fallback на factQuarter+factMoney считал плановые строки как фактические. '
     'Фикс: только factDatePresent (столбец Q).'),
    ('Plan count inflation', 'year.planCount++ без проверки planQ. '
     'Фикс: обёрнут в if(planQ).'),
    ('Budget key mismatch', 'summaryByPeriod использовал total-department ключи вместо KP-specific. '
     'Фикс: prefix.kp.qk.plan_fb.'),
    ('Reconciliation double-count', 'Official side считал только КП, calculated - КП+ЕП. '
     'Фикс: суммирование обоих методов на обеих сторонах.'),
]
for title, desc in bugs:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('12.3. Преимущества кода АЕМР перед СВОД', level=2)
advantages = [
    'Строжайшая обработка заглушек: 7 паттернов (Х, X, -, -, -, н/д, нет, не определена) vs 2 (Х, X)',
    'Безопасное деление: pct() возвращает 0 при нулевом знаменателе вместо #DIV/0!',
    'Fallback итогов: planTotalFor/factTotalFor используют сумму компонент при пустом итоге',
    'Детекция конфликтов: обнаружение расхождений между флагом AD и фактической экономией',
    'Явный перечень методов: COMPETITIVE_METHODS = {ЭА, ЭК, ЭЗК} точнее, чем L<>"ЕП"',
    'Классификация строк: score-based фильтр отсекает служебные/пустые строки, которые СВОД может включить',
]
for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_heading('12.4. Рекомендации', level=2)
recs = [
    ('Экономия', 'Внедрить гибридный метод: математический расчёт (J-X>0) как основа, '
     'флаг AD как подтверждение. При конфликте - маркировать для ручной проверки.'),
    ('СВОД заглушки', 'Рекомендовать администраторам СВОД дополнить COUNTIFS условиями Q<>"-" и Q<>"-" '
     'для полноты фильтрации фактов.'),
    ('ШДЮ перекрёстная сверка', 'Регулярно сверять сумму 3 месяцев из ШДЮ с квартальными значениями СВОД. '
     'Расхождение > 1% требует расследования.'),
    ('Trust score', 'После устранения дельт по экономии trust score должен достичь уровня A (>90). '
     'Текущий уровень B (76-83) обусловлен остаточными расхождениями.'),
]
for title, desc in recs:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('--- Конец документа ---')
r.italic = True

# Save
output_path = os.path.join(r'C:\Users\filat\dash', 'АЕМР_Аудит_Методик_Расчёта.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
