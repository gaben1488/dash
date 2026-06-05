"""
docx_to_text.py — конвертирует .docx → .txt (параграфы + таблицы) для чтения
агентами. Цель: 9 паспортов проекта (decision_engine, карта_маппинга,
карта_рисков, модуль_44фз, антикоррупционный_модуль и т.д.) — это
задокументированная спецификация системы.

Usage:
  python scripts/docx_to_text.py <dir|file> --out-dir <dir> [--recursive]
"""
import sys, os, glob, re
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

args = sys.argv[1:]
if not args:
    print("usage: docx_to_text.py <dir|file> --out-dir <dir> [--recursive]"); sys.exit(1)
target = args[0]
out_dir = args[args.index('--out-dir')+1] if '--out-dir' in args else '.'
recursive = '--recursive' in args
os.makedirs(out_dir, exist_ok=True)

if os.path.isdir(target):
    pat = os.path.join(target, '**', '*.docx') if recursive else os.path.join(target, '*.docx')
    files = sorted(glob.glob(pat, recursive=recursive))
    target_root = Path(target).resolve()
else:
    files = [target]
    target_root = None
files = [f for f in files if not os.path.basename(f).startswith('~$')]
print(f"[docx_to_text] {len(files)} файлов → {out_dir}")

def doc_text(path):
    d = Document(path)
    out = []
    for p in d.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for ti, t in enumerate(d.tables):
        out.append(f'\n[ТАБЛИЦА {ti+1}]')
        for r in t.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
            out.append(' | '.join(cells))
    return '\n'.join(out)

def output_name(path):
    p = Path(path).resolve()
    if recursive and target_root is not None:
        rel = p.relative_to(target_root).with_suffix('')
        raw = '__'.join(rel.parts)
    else:
        raw = p.stem
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', raw).strip(' ._')
    return (safe or 'document')[:180]

seen = {}
for f in files:
    safe = output_name(f)
    seen[safe] = seen.get(safe, 0) + 1
    if seen[safe] > 1:
        safe = f"{safe}_{seen[safe]}"
    try:
        txt = doc_text(f)
    except Exception as e:
        txt = f'[ОШИБКА чтения {f}: {e}]'
    Path(os.path.join(out_dir, safe + '.txt')).write_text(txt, encoding='utf-8')
    print(f'  {safe}.txt ({len(txt)} chars)')
print('[OK]')
